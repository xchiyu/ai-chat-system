import os
from typing import List, Optional
from llama_index.core import Document, VectorStoreIndex, StorageContext, load_index_from_storage
from llama_index.core.text_splitter import SentenceSplitter


def load_documents_from_folder(folder_path: str) -> List[Document]:
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                print(f"读取文本文件: {filename}，字符数: {len(text)}")
                documents.append(Document(text=text, metadata={"file_name": filename}))
            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_texts.append(cell.text.strip())
                full_text = "\n".join(paragraphs + table_texts)
                print(f"读取 Word 文件: {filename}，字符数: {len(full_text)}")
                documents.append(Document(text=full_text, metadata={"file_name": filename}))
            else:
                print(f"忽略不支持的文件: {filename}")
        except Exception as e:
            print(f"读取文件 {filename} 时出错: {e}")
            raise
    print(f"总共加载了 {len(documents)} 个文档。")
    return documents


def get_all_documents_content(folder_path: str) -> str:
    all_content = []
    for filename in sorted(os.listdir(folder_path)):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in [".txt", ".docx"]:
            continue
        
        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_texts.append(cell.text.strip())
                text = "\n".join(paragraphs + table_texts)
            
            name_without_ext = os.path.splitext(filename)[0]
            all_content.append(f"=== 文件: {name_without_ext} ===\n{text}")
        except Exception as e:
            print(f"读取文件 {filename} 时出错: {e}")
    
    return "\n\n".join(all_content)


def build_or_load_index(
    vector_path: str,
    docs_path: str,
    embed_model,
    chunk_size: int = 512,
    chunk_overlap: int = 50
):
    os.makedirs(vector_path, exist_ok=True)
    if not os.listdir(vector_path):
        print("向量库为空，开始构建...")
        docs = load_documents_from_folder(docs_path)
        if not docs:
            raise RuntimeError("未找到任何可读文档，请检查 docs/ 目录。")
        splitter = SentenceSplitter(
            separator="\n\n",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        nodes = splitter.get_nodes_from_documents(docs)
        index = VectorStoreIndex(nodes, embed_model=embed_model)
        index.storage_context.persist(vector_path)
        print("向量存储完成，已保存到", vector_path)
    else:
        print("检测到已有向量库，直接加载...")
        storage_context = StorageContext.from_defaults(persist_dir=vector_path)
        index = load_index_from_storage(storage_context, embed_model=embed_model)
        print("向量库加载完成")
    return index


class KnowledgeBase:
    def __init__(self, vector_path: str, docs_path: str, embed_model):
        self.vector_path = vector_path
        self.docs_path = docs_path
        self.embed_model = embed_model
        self.index = build_or_load_index(vector_path, docs_path, embed_model)
        self.retriever = self.index.as_retriever(similarity_top_k=5)

    def query(self, query_text: str, llm) -> str:
        query_engine = self.index.as_query_engine(llm=llm, similarity_top_k=5)
        response = query_engine.query(query_text)
        return response.response

    def get_relevant_context(self, query_text: str) -> str:
        nodes = self.retriever.retrieve(query_text)
        context = "\n\n".join([node.text for node in nodes])
        return context

    def get_all_documents_content(self) -> str:
        return get_all_documents_content(self.docs_path)

    def get_document_names(self) -> List[str]:
        names = []
        for filename in os.listdir(self.docs_path):
            file_path = os.path.join(self.docs_path, filename)
            if os.path.isfile(file_path):
                ext = os.path.splitext(filename)[1].lower()
                if ext in [".txt", ".docx"]:
                    names.append(os.path.splitext(filename)[0])
        return sorted(names)

    def rebuild_index(self):
        for file in os.listdir(self.vector_path):
            os.remove(os.path.join(self.vector_path, file))
        self.index = build_or_load_index(self.vector_path, self.docs_path, self.embed_model)
        self.retriever = self.index.as_retriever(similarity_top_k=5)
        print("向量库已重建")
