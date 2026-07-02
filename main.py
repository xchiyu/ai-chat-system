# main.py
import os
import sys
from llama_index.core import Document, VectorStoreIndex, StorageContext, load_index_from_storage
from llama_index.core.text_splitter import SentenceSplitter

from core.config import load_config
from core.models import get_zhipu_embedding, get_zhipu_llm

# ---------- 文档加载函数 ----------
def load_documents_from_folder(folder_path: str):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                print(f"读取文本文件: {filename}，字符数: {len(text)}")
                documents.append(Document(text=text, metadata={"file_name": filename}))
            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_texts.append(cell.text.strip())
                full_text = "\n".join(paragraphs + table_texts)
                print(f"读取 Word 文件: {filename}，字符数: {len(full_text)}")
                documents.append(Document(text=full_text, metadata={"file_name": filename}))
            else:
                print(f"⏭忽略不支持的文件: {filename}")
        except Exception as e:
            print(f"读取文件 {filename} 时出错: {e}")
            raise
    print(f"总共加载了 {len(documents)} 个文档。")
    return documents

# ---------- 构建或加载索引 ----------
def build_or_load_index(vector_path, docs_path, embed_model, chunk_size=512, chunk_overlap=50):
    os.makedirs(vector_path, exist_ok=True)
    if not os.listdir(vector_path):
        print("向量库为空，开始构建...")
        docs = load_documents_from_folder(docs_path)
        if not docs:
            raise RuntimeError("未找到任何可读文档，请检查 docs/ 目录。")
        splitter = SentenceSplitter(
            separator="\n\n",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        nodes = splitter.get_nodes_from_documents(docs)
        index = VectorStoreIndex(nodes, embed_model=embed_model)
        index.storage_context.persist(vector_path)
        print("向量存储完成，已保存到", vector_path)
    else:
        print("检测到已有向量库，直接加载...")
        storage_context = StorageContext.from_defaults(persist_dir=vector_path)
        index = load_index_from_storage(storage_context, embed_model=embed_model)
        print("向量库加载完成")
    return index

# ---------- 主程序 ----------
config = load_config("config.yaml")
embed_model = get_zhipu_embedding(config)
llm = get_zhipu_llm(config)

docs_path = "./docs/"
vector_path = "./vector_store/"

index = build_or_load_index(vector_path, docs_path, embed_model)
query_engine = index.as_query_engine(llm=llm, similarity_top_k=5)

print("\n知识库已就绪，输入 'exit' 退出。")
while True:
    try:
        query = input("\n请输入你的问题: ")
        if query.lower() == 'exit':
            break
        if not query.strip():
            continue
        response = query_engine.query(query)
        print(f"回答: {response.response}")
    except KeyboardInterrupt:
        print("\n退出程序")
        sys.exit(0)
    except Exception as e:
        print(f"发生错误: {e}")